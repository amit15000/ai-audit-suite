"""Simple test script for External Fact Check feature."""
from __future__ import annotations

import asyncio
import sys
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from app.core.config import get_settings
from app.services.comparison.hallucination.external_fact_check import (
    ClaimExtractor,
    ExternalFactCheckScorer,
)


async def test_claim_extraction():
    """Test claim extraction (both LLM and rule-based)."""
    print("=" * 60)
    print("TEST 1: Claim Extraction")
    print("=" * 60)
    
    settings = get_settings().external_fact_check
    
    # Test with LLM (default)
    print("\n1. Testing LLM-based extraction (default):")
    extractor_llm = ClaimExtractor(use_llm=True)
    
    response = """
    New York City, often referred to as the Big Apple, is the most populous city in the United States. According to the 2020 United States Census, the city has a population of approximately 8.8 million people, making it one of the largest metropolitan areas in the world. The city spans a total area of 302.6 square miles, which includes its five distinct boroughs: Manhattan, Brooklyn, Queens, The Bronx, and Staten Island.
    
    The city was originally founded in 1624 by Dutch settlers who established a trading post on the southern tip of Manhattan Island. This settlement was initially named New Amsterdam, after the Dutch capital. However, in 1664, the English captured the colony and renamed it New York in honor of the Duke of York, who later became King James II of England.
    
    One of the most iconic landmarks in New York City is the Statue of Liberty, which stands on Liberty Island in New York Harbor. This magnificent copper statue, a gift from France to the United States, was dedicated on October 28, 1886. Designed by French sculptor Frédéric Auguste Bartholdi, the statue stands 305 feet tall from the base to the tip of the torch and has become a universal symbol of freedom and democracy.
    
    New York City is also home to Central Park, one of the most famous urban parks in the world. The park covers 843 acres in the heart of Manhattan and was designed by Frederick Law Olmsted and Calvert Vaux. It was officially opened to the public in 1858 and has since become a vital green space for millions of residents and tourists.
    
    The city's economy is one of the largest in the world, with a gross metropolitan product of over $1.5 trillion as of 2023. Wall Street, located in Lower Manhattan, is considered the financial capital of the world and is home to the New York Stock Exchange, the largest stock exchange by market capitalization.
    
    According to a comprehensive 2022 study by the Metropolitan Transportation Authority, approximately 68% of New York City residents use public transportation for their daily commute, making it one of the highest public transit usage rates in the United States. The city's subway system, operated by the MTA, is one of the oldest and most extensive rapid transit systems in the world, with 472 stations across 27 lines.
    
    The Empire State Building, completed in 1931, was the world's tallest building for nearly 40 years until the completion of the World Trade Center's North Tower in 1970. Standing at 1,454 feet tall (including its antenna), it remains one of the most recognizable skyscrapers in the world and attracts millions of visitors annually.
    
    New York City is also renowned for its cultural institutions. The Metropolitan Museum of Art, founded in 1870, houses over 2 million works of art spanning 5,000 years of world culture. The museum's collection includes everything from ancient Egyptian artifacts to contemporary art, making it one of the largest and most comprehensive art museums in the world.
    """
    
    try:
        claims = await extractor_llm.extract_claims(response, max_claims=10)
        print(f"   [OK] Extracted {len(claims)} claims using LLM")
        for i, claim in enumerate(claims[:3], 1):
            print(f"   {i}. {claim.claim[:70]}...")
    except Exception as e:
        print(f"   [ERROR] LLM extraction failed: {e}")
        print("   → Falling back to rule-based...")
    
    # Test with rule-based fallback
    print("\n2. Testing rule-based extraction (fallback):")
    extractor_rule = ClaimExtractor(use_llm=False)
    claims = await extractor_rule.extract_claims(response, max_claims=10)
    print(f"   [OK] Extracted {len(claims)} claims using rule-based method")
    for i, claim in enumerate(claims[:3], 1):
        print(f"   {i}. {claim.claim[:70]}...")


async def test_full_scoring():
    """Test full external fact check scoring."""
    print("\n" + "=" * 60)
    print("TEST 2: Full External Fact Check Scoring (LLM + Web Search)")
    print("=" * 60)
    
    settings = get_settings().external_fact_check
    
    if not settings.enabled:
        print("\n⚠ External Fact Check is disabled in config.")
        print("   Set EXTERNAL_FACT_CHECK_ENABLED=true in .env")
        return
    
    # Check for OpenAI API key
    import os
    if not os.getenv("OPENAI_API_KEY"):
        print("\n[WARNING] OPENAI_API_KEY is not set in environment variables!")
        print("   This is required for LLM-based fact checking.")
        print("   Set OPENAI_API_KEY in your .env file or environment.")
        return
    
    print("\n[OK] Using LLM-based fact checking with OpenAI")
    print("   Watch the logs below for detailed claim extraction, evidence retrieval, and verification...\n")
    
    scorer = ExternalFactCheckScorer()
    
    # Single test response with multiple claims
    test_response = """
    New York City, often referred to as the Big Apple, is the most populous city in the United States. According to the 2020 United States Census, the city has a population of approximately 8.8 million people, making it one of the largest metropolitan areas in the world. The city spans a total area of 302.6 square miles, which includes its five distinct boroughs: Manhattan, Brooklyn, Queens, The Bronx, and Staten Island.
    
    The city was originally founded in 1624 by Dutch settlers who established a trading post on the southern tip of Manhattan Island. This settlement was initially named New Amsterdam, after the Dutch capital. However, in 1664, the English captured the colony and renamed it New York in honor of the Duke of York, who later became King James II of England.
    
    One of the most iconic landmarks in New York City is the Statue of Liberty, which stands on Liberty Island in New York Harbor. This magnificent copper statue, a gift from France to the United States, was dedicated on October 28, 1886. Designed by French sculptor Frédéric Auguste Bartholdi, the statue stands 305 feet tall from the base to the tip of the torch and has become a universal symbol of freedom and democracy.
    
    New York City is also home to Central Park, one of the most famous urban parks in the world. The park covers 843 acres in the heart of Manhattan and was designed by Frederick Law Olmsted and Calvert Vaux. It was officially opened to the public in 1858 and has since become a vital green space for millions of residents and tourists.
    
    The city's economy is one of the largest in the world, with a gross metropolitan product of over $1.5 trillion as of 2023. Wall Street, located in Lower Manhattan, is considered the financial capital of the world and is home to the New York Stock Exchange, the largest stock exchange by market capitalization.
    
    According to a comprehensive 2022 study by the Metropolitan Transportation Authority, approximately 68% of New York City residents use public transportation for their daily commute, making it one of the highest public transit usage rates in the United States. The city's subway system, operated by the MTA, is one of the oldest and most extensive rapid transit systems in the world, with 472 stations across 27 lines.
    
    The Empire State Building, completed in 1931, was the world's tallest building for nearly 40 years until the completion of the World Trade Center's North Tower in 1970. Standing at 1,454 feet tall (including its antenna), it remains one of the most recognizable skyscrapers in the world and attracts millions of visitors annually.
    
    New York City is also renowned for its cultural institutions. The Metropolitan Museum of Art, founded in 1870, houses over 2 million works of art spanning 5,000 years of world culture. The museum's collection includes everything from ancient Egyptian artifacts to contemporary art, making it one of the largest and most comprehensive art museums in the world.
    """
    
    print(f"\nTesting Response:")
    print(f"   {test_response[:100]}...")
    
    try:
        result = await scorer.calculate_sub_score(test_response)
        
        print(f"\n{'='*60}")
        print(f"   RESULTS SUMMARY")
        print(f"{'='*60}")
        print(f"   • Score: {result.score}/100")
        print(f"   • Coverage: {result.coverage:.1%}")
        print(f"   • Claims Verified: {len(result.claims)}")
        print(f"   • Total Sources Used: {len(result.sources_used)}")
        
        if result.claims:
            print(f"\n{'='*60}")
            print(f"   CLAIMS LIST ({len(result.claims)}):")
            print(f"{'='*60}")
            
            for i, claim in enumerate(result.claims, 1):
                print(f"\n   {i}. {claim.claim}")
                
                # Verification result
                is_true = claim.verdict == "SUPPORTED"
                status_icon = "[TRUE]" if is_true else "[FALSE]"
                status_text = "TRUE" if is_true else "FALSE"
                print(f"      {status_icon} Result: {status_text}")
                
                # Sources (full URLs from OpenAI)
                if claim.top_evidence:
                    print(f"      Sources ({len(claim.top_evidence)}):")
                    for j, evidence in enumerate(claim.top_evidence, 1):
                        # Show full URL if available, otherwise domain
                        source_display = evidence.url if evidence.url and evidence.url.startswith("http") else evidence.domain
                        print(f"         • {source_display}")
                else:
                    print(f"      Sources: None provided")
                
                # Explanation (from evidence snippet which contains explanation)
                if claim.top_evidence and claim.top_evidence[0].snippet:
                    explanation = claim.top_evidence[0].snippet
                    if len(explanation) > 150:
                        explanation = explanation[:150] + "..."
                    print(f"      Explanation: {explanation}")
        
        if result.notes:
            print(f"\n   Notes: {', '.join(result.notes)}")
        
        # Generate .docx report (one-time)
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            doc = Document()
            
            # Title
            title = doc.add_heading('External Fact Check Results Summary', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(10)
            date_para.runs[0].font.italic = True
            
            doc.add_paragraph()  # Spacing
            
            # Summary Section
            doc.add_heading('Summary', 1)
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Score: ").bold = True
            summary_para.add_run(f"{result.score}/100")
            summary_para.add_run(f"\nCoverage: ").bold = True
            summary_para.add_run(f"{result.coverage:.1%}")
            summary_para.add_run(f"\nClaims Verified: ").bold = True
            summary_para.add_run(f"{len(result.claims)}")
            summary_para.add_run(f"\nTotal Sources Used: ").bold = True
            summary_para.add_run(f"{len(result.sources_used)}")
            
            doc.add_paragraph()  # Spacing
            
            # Claims Section
            doc.add_heading('Claims Verification Details', 1)
            
            for i, claim in enumerate(result.claims, 1):
                # Claim number and text
                claim_heading = doc.add_heading(f"Claim {i}", 2)
                claim_para = doc.add_paragraph(claim.claim)
                claim_para.style = 'List Paragraph'
                
                # Verification result
                result_para = doc.add_paragraph()
                is_true = claim.verdict == "SUPPORTED"
                status_text = "TRUE" if is_true else "FALSE"
                result_para.add_run("Result: ").bold = True
                status_run = result_para.add_run(status_text)
                status_run.bold = True
                status_run.font.color.rgb = RGBColor(0, 128, 0) if is_true else RGBColor(255, 0, 0)
                
                # Sources
                if claim.top_evidence:
                    sources_para = doc.add_paragraph()
                    sources_para.add_run(f"Sources ({len(claim.top_evidence)}):").bold = True
                    for evidence in claim.top_evidence:
                        source_url = evidence.url if evidence.url and evidence.url.startswith("http") else f"https://{evidence.domain}"
                        source_item = doc.add_paragraph(source_url, style='List Bullet')
                        # Make URL clickable (add hyperlink)
                        try:
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            hyperlink = source_item._element.get_or_add_hyperlink(source_url)
                            run = source_item.runs[0] if source_item.runs else source_item.add_run(source_url)
                            run.font.color.rgb = RGBColor(0, 0, 255)
                            run.font.underline = True
                        except:
                            pass  # If hyperlink fails, just show text
                else:
                    no_sources = doc.add_paragraph()
                    no_sources.add_run("Sources: ").bold = True
                    no_sources.add_run("None provided")
                
                # Explanation
                if claim.top_evidence and claim.top_evidence[0].snippet:
                    explanation = claim.top_evidence[0].snippet
                    expl_para = doc.add_paragraph()
                    expl_para.add_run("Explanation: ").bold = True
                    expl_para.add_run(explanation)
                
                doc.add_paragraph()  # Spacing between claims
            
            # Notes
            if result.notes:
                doc.add_heading('Notes', 1)
                notes_para = doc.add_paragraph(', '.join(result.notes))
            
            # Save document
            output_path = project_root / "external_fact_check_results.docx"
            doc.save(str(output_path))
            print(f"\n{'='*60}")
            print(f"   DOCX REPORT GENERATED")
            print(f"{'='*60}")
            print(f"   File saved to: {output_path}")
            print(f"{'='*60}")
            
        except ImportError:
            print(f"\n[INFO] python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            print(f"\n[WARNING] Could not generate .docx file: {e}")
            
    except Exception as e:
        print(f"   [ERROR] Error: {e}")
        import traceback
        traceback.print_exc()


async def test_configuration():
    """Test configuration settings."""
    print("\n" + "=" * 60)
    print("TEST 3: Configuration Check")
    print("=" * 60)
    
    settings = get_settings().external_fact_check
    
    import os
    
    print(f"\nConfiguration:")
    print(f"  • Enabled: {settings.enabled}")
    print(f"  • OpenAI API Key: {'[SET]' if os.getenv('OPENAI_API_KEY') else '[NOT SET - REQUIRED]'}")
    print(f"  • Use LLM for extraction: {settings.claim_extraction_use_llm}")
    print(f"  • Max Claims: {settings.max_claims_per_response}")
    print(f"  • Verification Timeout: {settings.verification_timeout}s")
    
    if not settings.enabled:
        print("\n[WARNING] External Fact Check is disabled!")
    print(f"  • Search Provider: {settings.search_provider}")
    print(f"  • Top K Results: {settings.top_k_results}")
    print(f"  • Search Timeout: {settings.search_timeout}s")


async def main():
    """Run all tests."""
    print("\n" + "=" * 60)
    print("EXTERNAL FACT CHECK TEST SUITE")
    print("=" * 60)
    
    # Check configuration first
    await test_configuration()
    
    # Test claim extraction
    try:
        await test_claim_extraction()
    except Exception as e:
        print(f"\n[ERROR] Claim extraction test failed: {e}")
        import traceback
        traceback.print_exc()
    
    # Test full scoring (may take longer due to API calls)
    print("\n" + "=" * 60)
    print("Running full scoring test (this may take 30-60 seconds)...")
    print("=" * 60)
    
    try:
        await test_full_scoring()
    except Exception as e:
        print(f"\n[ERROR] Full scoring test failed: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n" + "=" * 60)
    print("TEST SUITE COMPLETE")
    print("=" * 60)
    print("\nFor more detailed testing, see:")
    print("  • Unit tests: pytest tests/test_external_fact_check.py -v")
    print("  • Documentation: docs/EXTERNAL_FACT_CHECK_TESTING.md")


if __name__ == "__main__":
    asyncio.run(main())
