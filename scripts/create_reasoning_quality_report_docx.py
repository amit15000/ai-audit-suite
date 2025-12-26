"""One-time script to create DOCX report for Reasoning Quality scoring rules and test results."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

project_root = Path(__file__).parent.parent

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Reasoning Quality Scoring System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(10)
date_para.runs[0].font.italic = True

doc.add_paragraph()  # Spacing

# ============================================================================
# SECTION 1: SCORING RULES & METHODOLOGY
# ============================================================================
doc.add_heading('1. Scoring Rules & Methodology', level=1)

doc.add_paragraph(
    'The Reasoning Quality scoring system evaluates the logical structure, coherence, and validity '
    'of reasoning in AI-generated responses. It uses a weighted average of four sub-scores, each '
    'measuring different aspects of reasoning quality. The system employs LLM-based semantic analysis '
    'to assess reasoning patterns, detect logical gaps, and identify fallacious reasoning.'
)

doc.add_paragraph()  # Spacing

# Overall Score Formula Section
doc.add_heading('1.1 Overall Score Calculation', level=2)
doc.add_paragraph('The overall reasoning quality score is calculated using a weighted average:')
formula_para = doc.add_paragraph()
formula_para.add_run('Overall Score = (Step-by-Step × 0.25) + (Logical Consistency × 0.30) + (Missing Steps × 0.25) + (Wrong Logic × 0.20)\n\n').bold = True
formula_para.add_run('Where each sub-score ranges from 0-10, and weights sum to 1.0 (100%)')

doc.add_paragraph()  # Spacing

# Weight Distribution Table
doc.add_heading('1.2 Weight Distribution', level=2)
doc.add_paragraph('Each sub-score contributes to the overall score with the following weights:')

table_weights = doc.add_table(rows=5, cols=2)
table_weights.style = 'Light Grid Accent 1'

# Header row
header_cells = table_weights.rows[0].cells
header_cells[0].text = 'Sub-Score Component'
header_cells[1].text = 'Weight (Percentage)'
for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)

# Data rows
table_weights.rows[1].cells[0].text = '1. Step-by-Step Reasoning'
table_weights.rows[1].cells[1].text = '25%'
table_weights.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
table_weights.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True

table_weights.rows[2].cells[0].text = '2. Logical Consistency'
table_weights.rows[2].cells[1].text = '30%'
table_weights.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
table_weights.rows[2].cells[1].paragraphs[0].runs[0].font.bold = True

table_weights.rows[3].cells[0].text = '3. Missing Steps Detection'
table_weights.rows[3].cells[1].text = '25%'
table_weights.rows[3].cells[0].paragraphs[0].runs[0].font.bold = True
table_weights.rows[3].cells[1].paragraphs[0].runs[0].font.bold = True

table_weights.rows[4].cells[0].text = '4. Wrong Logic Detection'
table_weights.rows[4].cells[1].text = '20%'
table_weights.rows[4].cells[0].paragraphs[0].runs[0].font.bold = True
table_weights.rows[4].cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()  # Spacing

# Sub-Score Descriptions
doc.add_heading('1.3 Sub-Score Components', level=2)

doc.add_heading('1.3.1 Step-by-Step Reasoning (25% weight)', level=3)
doc.add_paragraph('Measures the clarity and structure of logical progression:')
step_by_step_para = doc.add_paragraph()
step_by_step_para.add_run('• ').bold = True
step_by_step_para.add_run('Clear enumeration and organization of ideas\n')
step_by_step_para.add_run('• ').bold = True
step_by_step_para.add_run('Explicit logical connectors and transitions\n')
step_by_step_para.add_run('• ').bold = True
step_by_step_para.add_run('Systematic breakdown of arguments\n')
step_by_step_para.add_run('• ').bold = True
step_by_step_para.add_run('Logical flow from premises to conclusions\n')
step_by_step_para.add_run('• ').bold = True
step_by_step_para.add_run('Score Range: 0-10 (10 = excellent structure, 0 = no structure)')

doc.add_paragraph()  # Spacing

doc.add_heading('1.3.2 Logical Consistency (30% weight)', level=3)
doc.add_paragraph('Measures internal consistency and absence of contradictions:')
consistency_para = doc.add_paragraph()
consistency_para.add_run('• ').bold = True
consistency_para.add_run('Reuses contradiction detection from Hallucination Score\n')
consistency_para.add_run('• ').bold = True
consistency_para.add_run('Detects semantic contradictions in statements\n')
consistency_para.add_run('• ').bold = True
consistency_para.add_run('Identifies incompatible claims within the response\n')
consistency_para.add_run('• ').bold = True
consistency_para.add_run('Score Range: 0-10 (10 = fully consistent, 0 = many contradictions)')
doc.add_paragraph()
doc.add_paragraph('Note: This component uses the contradictory information score from the Hallucination Score system, ensuring consistency across the audit framework.').italic = True

doc.add_paragraph()  # Spacing

doc.add_heading('1.3.3 Missing Steps Detection (25% weight)', level=3)
doc.add_paragraph('Measures completeness of reasoning chains:')
missing_steps_para = doc.add_paragraph()
missing_steps_para.add_run('• ').bold = True
missing_steps_para.add_run('Identifies gaps in logical reasoning\n')
missing_steps_para.add_run('• ').bold = True
missing_steps_para.add_run('Detects abrupt jumps from premises to conclusions\n')
missing_steps_para.add_run('• ').bold = True
missing_steps_para.add_run('Evaluates whether intermediate steps are explained\n')
missing_steps_para.add_run('• ').bold = True
missing_steps_para.add_run('Score Range: 0-10 (10 = complete reasoning chain, 0 = major gaps)')

doc.add_paragraph()  # Spacing

doc.add_heading('1.3.4 Wrong Logic Detection (20% weight)', level=3)
doc.add_paragraph('Measures the presence of logical fallacies and invalid reasoning:')
wrong_logic_para = doc.add_paragraph()
wrong_logic_para.add_run('• ').bold = True
wrong_logic_para.add_run('Detects common logical fallacies (ad hominem, slippery slope, etc.)\n')
wrong_logic_para.add_run('• ').bold = True
wrong_logic_para.add_run('Identifies invalid reasoning patterns\n')
wrong_logic_para.add_run('• ').bold = True
wrong_logic_para.add_run('Evaluates whether conclusions follow from premises\n')
wrong_logic_para.add_run('• ').bold = True
wrong_logic_para.add_run('Score Range: 0-10 (10 = no fallacies, 0 = many fallacies)')

doc.add_paragraph()  # Spacing

# Example Calculation
doc.add_heading('1.4 Example Calculation', level=2)

example_para = doc.add_paragraph()
example_para.add_run('Example: Response with the following sub-scores:\n\n').bold = True
example_para.add_run('• Step-by-Step Reasoning: 9/10\n')
example_para.add_run('• Logical Consistency: 8/10\n')
example_para.add_run('• Missing Steps Detection: 6/10\n')
example_para.add_run('• Wrong Logic Detection: 10/10\n\n')
example_para.add_run('Calculation:\n').bold = True
example_para.add_run('Overall Score = (9 × 0.25) + (8 × 0.30) + (6 × 0.25) + (10 × 0.20)\n')
example_para.add_run('            = 2.25 + 2.40 + 1.50 + 2.00\n')
example_para.add_run('            = 8.15/10').bold = True

doc.add_page_break()

# ============================================================================
# SECTION 2: TEST RESULTS
# ============================================================================
doc.add_heading('2. Comprehensive Test Results', level=1)

doc.add_paragraph(
    'The following sections contain detailed test results demonstrating the Reasoning Quality '
    'scoring system across various scenarios, from excellent reasoning to poor logical structure.'
)

doc.add_paragraph()  # Spacing

# TEST 1
doc.add_heading('TEST 1: Excellent Reasoning - Well-Structured Analysis', level=2)

response1 = """To understand why renewable energy is crucial for our future, we need to examine three key factors.

First, we must consider environmental impact. The burning of fossil fuels releases carbon dioxide 
into the atmosphere, contributing to a 1.1°C increase in global average temperatures since pre-industrial times. 
This temperature rise causes ice sheet melting, sea level rise, and extreme weather patterns.

Second, we must examine economic viability. Historically, renewable energy was expensive, but 
technological advances have dramatically reduced costs. Solar panel prices have dropped by 90% 
over the past decade. Wind energy is now cost-competitive with fossil fuels in many regions. 
These cost reductions make renewable energy economically attractive.

Third, let's consider resource availability. Fossil fuels are finite - coal, oil, and natural 
gas reserves will eventually be depleted. In contrast, renewable sources like sunlight and wind 
are essentially limitless and available worldwide.

Therefore, transitioning to renewable energy addresses environmental concerns, makes economic 
sense, and ensures long-term energy security. This combination of factors makes renewable energy 
not just desirable, but necessary for sustainable development."""

doc.add_paragraph('Description: Response with excellent step-by-step reasoning, complete logical chain, and sound logic').italic = True
doc.add_paragraph('Response:').bold = True
doc.add_paragraph(response1, style='Intense Quote')

doc.add_paragraph('Results:').bold = True
results1 = doc.add_paragraph()
results1.add_run('1️⃣  STEP-BY-STEP REASONING\n').bold = True
results1.add_run('   Score: 9/10\n')
results1.add_run('   Explanation: ').bold = True
results1.add_run(
    'The response demonstrates a strong step-by-step reasoning quality with a clear structure and logical progression. '
    'The author enumerates three key factors (environmental impact, economic viability, and resource availability), '
    'which organizes the argument effectively. Each factor is explicitly stated and developed with supporting evidence, '
    'such as the data on carbon emissions and the cost reduction of renewable technologies. Furthermore, logical '
    'connectors like "therefore" and "first, second, third" facilitate smooth transitions between ideas. The reasoning '
    'flows logically from one premise to the next, culminating in the conclusion that renewable energy is necessary '
    'for sustainable development. However, the response could improve slightly by including more explicit logical '
    'connections between the implications of each factor and the final conclusion, which is why it does not receive a perfect score.\n\n'
)

results1.add_run('2️⃣  LOGICAL CONSISTENCY\n').bold = True
results1.add_run('   Score: 8/10\n')
results1.add_run('   Explanation: Reused from Hallucination Score (contradiction detection): 8/10\n\n')

results1.add_run('3️⃣  MISSING STEPS DETECTION\n').bold = True
results1.add_run('   Score: 6/10\n')
results1.add_run('   Explanation: ').bold = True
results1.add_run(
    'While the response provides some strong points regarding environmental impact, economic viability, and resource '
    'availability, it lacks explicit connections that directly tie these premises to the necessity of renewable energy. '
    'For instance, the argument could benefit from more detailed explanations of how each factor directly correlates to '
    'the "necessity for sustainable development." There are also abrupt transitions between points with insufficient '
    'linkage to emphasize how they collectively support the main conclusion.\n\n'
)

results1.add_run('4️⃣  WRONG LOGIC DETECTION\n').bold = True
results1.add_run('   Score: 10/10\n')
results1.add_run('   Explanation: ').bold = True
results1.add_run(
    'The argument is structured logically with valid reasoning. It presents well-supported premises regarding environmental '
    'impact, economic viability, and resource availability, all leading to a coherent conclusion that transitioning to '
    'renewable energy is necessary for sustainable development. No logical fallacies or invalid reasoning patterns are detected.\n\n'
)

results1.add_run('📊 OVERALL ASSESSMENT\n').bold = True
results1.add_run('   Weighted Average: 8.15/10\n')
results1.add_run('   (Step-by-Step: 25%, Consistency: 30%, Missing Steps: 25%, Wrong Logic: 20%)').bold = True

doc.add_paragraph()  # Spacing

# TEST 2
doc.add_heading('TEST 2: Missing Logical Steps - Jumping to Conclusions', level=2)

response2 = """Artificial intelligence is becoming more advanced. Machine learning models can now generate
text, images, and code. Therefore, AI will replace all human jobs within the next five years.
We need to prepare for this massive disruption. The solution is universal basic income and
retraining programs."""

doc.add_paragraph('Description: Response that jumps from premises to conclusions without proper intermediate reasoning').italic = True
doc.add_paragraph('Response:').bold = True
doc.add_paragraph(response2, style='Intense Quote')

doc.add_paragraph('Results:').bold = True
results2 = doc.add_paragraph()
results2.add_run('1️⃣  STEP-BY-STEP REASONING\n').bold = True
results2.add_run('   Score: 4/10\n')
results2.add_run('   Explanation: ').bold = True
results2.add_run(
    'The response presents a series of claims but lacks a clear structured progression and explicit reasoning steps. '
    'While it begins by stating that AI is becoming more advanced and can generate various outputs, it abruptly jumps '
    'to the conclusion that AI will replace all human jobs without providing supporting evidence or logical reasoning '
    'for this leap. The transition between ideas is minimal, with only "therefore" used as a connector, which does not '
    'sufficiently link the statements. There is no elaboration on why universal basic income and retraining programs '
    'are the proposed solutions, leading to disorganized reasoning and unclear argument flow.\n\n'
)

results2.add_run('2️⃣  LOGICAL CONSISTENCY\n').bold = True
results2.add_run('   Score: 8/10\n')
results2.add_run('   Explanation: Reused from Hallucination Score (contradiction detection): 8/10\n\n')

results2.add_run('3️⃣  MISSING STEPS DETECTION\n').bold = True
results2.add_run('   Score: 4/10\n')
results2.add_run('   Explanation: ').bold = True
results2.add_run(
    'The argument contains significant gaps in reasoning. There is an abrupt jump from the claim that AI is becoming '
    'more advanced to the conclusion that it will replace all human jobs within five years, lacking justification for '
    'this timeframe and the sweeping claim. The connection between AI advancements and job loss is not elaborated, '
    'leaving the conclusion unsupported. Additionally, the proposed solutions of universal basic income and retraining '
    'programs are introduced without explanation of how they would effectively address the job displacement caused by AI.\n\n'
)

results2.add_run('4️⃣  WRONG LOGIC DETECTION\n').bold = True
results2.add_run('   Score: 4/10\n')
results2.add_run('   Explanation: ').bold = True
results2.add_run(
    'The reasoning contains a false cause fallacy, implying that advancements in AI directly lead to the conclusion that '
    'all human jobs will be replaced within five years without sufficient evidence to support this drastic prediction. '
    'Additionally, the argument suffers from hasty generalization, as the claim about job replacement is based on a broad '
    'assumption about the future capabilities of technology without considering various factors that could influence employment. '
    'Overall, the conclusions do not follow logically from the premises provided.\n\n'
)

results2.add_run('📊 OVERALL ASSESSMENT\n').bold = True
results2.add_run('   Weighted Average: 5.20/10\n')
results2.add_run('   (Step-by-Step: 25%, Consistency: 30%, Missing Steps: 25%, Wrong Logic: 20%)').bold = True

doc.add_paragraph()  # Spacing

# TEST 3
doc.add_heading('TEST 3: Logical Fallacies - Multiple Reasoning Errors', level=2)

response3 = """Dr. Smith's research on climate change should be dismissed because she drives a gasoline-powered
car. If we implement carbon taxes, the government will control every aspect of our lives. Since
some people got sick after getting vaccinated, vaccines must be dangerous. All credible scientists
agree with my position because those who don't are obviously paid by big corporations. My argument
is correct because I'm an expert, and experts are always right about everything."""

doc.add_paragraph('Description: Response containing various logical fallacies and invalid reasoning patterns').italic = True
doc.add_paragraph('Response:').bold = True
doc.add_paragraph(response3, style='Intense Quote')

doc.add_paragraph('Results:').bold = True
results3 = doc.add_paragraph()
results3.add_run('1️⃣  STEP-BY-STEP REASONING\n').bold = True
results3.add_run('   Score: 1/10\n')
results3.add_run('   Explanation: ').bold = True
results3.add_run(
    'The response demonstrates very poor step-by-step reasoning quality. There is no clear structure or logical progression; '
    'the claims are disjointed and lack proper transitions. For example, the reasoning jumps abruptly from Dr. Smith\'s car '
    'to carbon taxes without connecting these ideas logically. There are no explicit logical connectors to clarify relationships '
    'between statements, and the argument lacks a coherent flow. Additionally, the assertions made about experts are not '
    'substantiated, further weakening the reasoning. Overall, the response presents random, disconnected ideas without effective '
    'argumentative structure.\n\n'
)

results3.add_run('2️⃣  LOGICAL CONSISTENCY\n').bold = True
results3.add_run('   Score: 8/10\n')
results3.add_run('   Explanation: Reused from Hallucination Score (contradiction detection): 8/10\n\n')

results3.add_run('3️⃣  MISSING STEPS DETECTION\n').bold = True
results3.add_run('   Score: 1/10\n')
results3.add_run('   Explanation: ').bold = True
results3.add_run(
    'The response contains critical gaps in logical reasoning. It abruptly transitions from Dr. Smith\'s car choice to '
    'dismissing her research without any rationale linking the two. The claim about carbon taxes controlling lives lacks '
    'supporting arguments that connect those two concepts. The assertion about vaccine danger based on anecdotal evidence '
    'is unsupported and jumps to a conclusion without proper reasoning. Additionally, the statement about all credible '
    'scientists agreeing with the position lacks evidence and fails to explain why dissenting scientists are discredited. '
    'Finally, claiming that expertise guarantees correctness is flawed and unsupported by logical steps, leading to critical '
    'points of ambiguity and incomplete reasoning.\n\n'
)

results3.add_run('4️⃣  WRONG LOGIC DETECTION\n').bold = True
results3.add_run('   Score: 1/10\n')
results3.add_run('   Explanation: ').bold = True
results3.add_run(
    'The response contains multiple logical fallacies including ad hominem (attacking Dr. Smith for driving a gasoline car '
    'instead of addressing her research), slippery slope (claiming carbon taxes will lead to total government control), '
    'hasty generalization (concluding vaccines are dangerous based on a few cases), appeal to authority (suggesting all '
    'credible scientists agree without evidence), and circular reasoning (arguing the author\'s expertise makes their argument '
    'correct). These fallacies indicate critical logical errors and extensive fallacious reasoning.\n\n'
)

results3.add_run('📊 OVERALL ASSESSMENT\n').bold = True
results3.add_run('   Weighted Average: 3.10/10\n')
results3.add_run('   (Step-by-Step: 25%, Consistency: 30%, Missing Steps: 25%, Wrong Logic: 20%)').bold = True

doc.add_paragraph()  # Spacing

# TEST 4
doc.add_heading('TEST 4: Inconsistent Information - Contradictory Claims', level=2)

response4 = """The company reported excellent financial performance this quarter. Revenue increased by 30%
compared to the same period last year, reaching $50 million. Our profit margins improved
significantly due to operational efficiency. However, we also experienced a 20% decline in
revenue this quarter, with total earnings dropping to $30 million. The CEO announced that
the company is struggling financially and may need to lay off employees. Despite these
challenges, we're confident about our strong financial position and continued growth."""

doc.add_paragraph('Description: Response with internal contradictions that make the reasoning inconsistent').italic = True
doc.add_paragraph('Response:').bold = True
doc.add_paragraph(response4, style='Intense Quote')

doc.add_paragraph('Results:').bold = True
results4 = doc.add_paragraph()
results4.add_run('1️⃣  STEP-BY-STEP REASONING\n').bold = True
results4.add_run('   Score: 4/10\n')
results4.add_run('   Explanation: ').bold = True
results4.add_run(
    'The response lacks clear step-by-step reasoning and has significant disorganization. There are contradictory statements '
    'about financial performance without a clear logical progression, such as first stating excellent performance followed by '
    'a decline in revenue. Transitions are weak or absent, leading to abrupt jumps in reasoning. For example, the statement '
    'about operational efficiency is not clearly connected to the subsequent decline in revenue. Overall, the argument does '
    'not flow logically from premises to conclusions, creating confusion.\n\n'
)

results4.add_run('2️⃣  LOGICAL CONSISTENCY\n').bold = True
results4.add_run('   Score: 2/10\n')
results4.add_run('   Explanation: Reused from Hallucination Score (contradiction detection): 2/10\n\n')

results4.add_run('3️⃣  MISSING STEPS DETECTION\n').bold = True
results4.add_run('   Score: 4/10\n')
results4.add_run('   Explanation: ').bold = True
results4.add_run(
    'The response presents conflicting information about financial performance, with a reported increase in revenue followed '
    'by a decline, which creates confusion. There is a lack of clarity regarding how the company can be financially strong '
    'while facing a revenue drop and possible layoffs. The transitions between claims about operational efficiency and the '
    'decline in revenue are abrupt and unsupported. Additionally, the conclusion of a strong financial position lacks sufficient '
    'reasoning or evidence, indicating that the overall reasoning chain is incomplete.\n\n'
)

results4.add_run('4️⃣  WRONG LOGIC DETECTION\n').bold = True
results4.add_run('   Score: 4/10\n')
results4.add_run('   Explanation: ').bold = True
results4.add_run(
    'The response presents conflicting financial information, stating both significant revenue growth of 30% and a subsequent '
    'decline of 20%, leading to confusion and inconsistency. The claim of being in a "strong financial position" is also '
    'misaligned with the CEO\'s statement about potential layoffs, creating a non sequitur where the conclusion does not '
    'logically follow from the premises. Additionally, the argument lacks clarity and coherence, which undermines its overall validity.\n\n'
)

results4.add_run('📊 OVERALL ASSESSMENT\n').bold = True
results4.add_run('   Weighted Average: 3.40/10\n')
results4.add_run('   (Step-by-Step: 25%, Consistency: 30%, Missing Steps: 25%, Wrong Logic: 20%)').bold = True

doc.add_paragraph()  # Spacing

# TEST 5
doc.add_heading('TEST 5: Unstructured Reasoning - Poor Organization', level=2)

response5 = """There are many factors to consider. Technology changes things. Society evolves over time.
People have different opinions. Economic conditions matter. Some solutions work better than
others. Historical context is important. Cultural differences exist. Various approaches have
been tried. Results can vary. The situation is complex. Multiple perspectives should be
considered. Change happens gradually sometimes. Innovation drives progress. Collaboration
helps. Understanding comes from experience. Different methods suit different situations.
Overall, it's a multifaceted issue requiring careful analysis."""

doc.add_paragraph('Description: Response with valid information but poor structure and unclear reasoning flow').italic = True
doc.add_paragraph('Response:').bold = True
doc.add_paragraph(response5, style='Intense Quote')

doc.add_paragraph('Results:').bold = True
results5 = doc.add_paragraph()
results5.add_run('1️⃣  STEP-BY-STEP REASONING\n').bold = True
results5.add_run('   Score: 3/10\n')
results5.add_run('   Explanation: ').bold = True
results5.add_run(
    'The response lacks a clear, structured logical progression and does not explicitly state reasoning steps. While it lists '
    'several factors and observations, the ideas are presented in a disorganized manner without proper transitions or logical '
    'connectors. Phrases such as "many factors to consider" and "overall, it\'s a multifaceted issue" suggest complexity but '
    'do not link the individual points coherently to an argument. There is no systematic enumeration or breakdown, leading to '
    'an unclear flow of ideas and reasoning.\n\n'
)

results5.add_run('2️⃣  LOGICAL CONSISTENCY\n').bold = True
results5.add_run('   Score: 8/10\n')
results5.add_run('   Explanation: Reused from Hallucination Score (contradiction detection): 8/10\n\n')

results5.add_run('3️⃣  MISSING STEPS DETECTION\n').bold = True
results5.add_run('   Score: 3/10\n')
results5.add_run('   Explanation: ').bold = True
results5.add_run(
    'The response lacks a clear logical progression with abrupt transitions between ideas. It mentions various factors that '
    'influence a situation but fails to connect these points to a specific conclusion or analysis, resulting in unsupported '
    'assertions. Additionally, the reasoning chain is incomplete, as it does not provide any specific examples or detailing '
    'how these factors interact or lead to a conclusion. The overall argument is fragmented and does not follow a coherent flow.\n\n'
)

results5.add_run('4️⃣  WRONG LOGIC DETECTION\n').bold = True
results5.add_run('   Score: 3/10\n')
results5.add_run('   Explanation: ').bold = True
results5.add_run(
    'The response uses vague and general statements without any specific argument or evidence, leading to a lack of logical '
    'coherence. It presents a non sequitur by failing to connect the various factors mentioned to a clear conclusion or position. '
    'The repetitive nature of the statements indicates a circular reasoning pattern where the complexity of the issue is stated '
    'without addressing any specific argument or claim. Overall, it lacks a sound argument structure, making it difficult to '
    'follow or evaluate logically.\n\n'
)

results5.add_run('📊 OVERALL ASSESSMENT\n').bold = True
results5.add_run('   Weighted Average: 4.50/10\n')
results5.add_run('   (Step-by-Step: 25%, Consistency: 30%, Missing Steps: 25%, Wrong Logic: 20%)').bold = True

doc.add_paragraph()  # Spacing

# ============================================================================
# SUMMARY
# ============================================================================
doc.add_page_break()
doc.add_heading('Summary', level=1)

summary_para = doc.add_paragraph()
summary_para.add_run('The Reasoning Quality scoring system successfully:').bold = True
summary_para.add_run('\n• Evaluates logical structure and coherence using four weighted sub-scores')
summary_para.add_run('\n• Detects missing steps and logical gaps in reasoning chains')
summary_para.add_run('\n• Identifies logical fallacies and invalid reasoning patterns')
summary_para.add_run('\n• Measures internal consistency through contradiction detection')
summary_para.add_run('\n• Provides weighted average scoring that emphasizes logical consistency (30%)')
summary_para.add_run('\n• Distinguishes between well-structured arguments and poor reasoning')
summary_para.add_run('\n• Offers detailed explanations for each sub-score component')

# Save document
output_path = project_root / "Reasoning_Quality_Scoring_Report.docx"
doc.save(str(output_path))

print(f"\n{'='*80}")
print(f"   DOCX REPORT GENERATED SUCCESSFULLY")
print(f"{'='*80}")
print(f"   File saved to: {output_path}")
print(f"{'='*80}")


