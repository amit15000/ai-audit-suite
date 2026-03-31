"""Generate DOCX with scoring rules first, then complete test results."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

project_root = Path(__file__).parent.parent
doc = Document()

# Style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Code Vulnerability Auditor - Complete Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================================
# PART 1: ALL SCORING RULES
# ============================================================================
doc.add_heading('PART 1: SCORING RULES & MATHEMATICS', 0)

doc.add_heading('1. Overview', level=1)
doc.add_paragraph('The Code Vulnerability Auditor analyzes code for 5 categories: Security Flaws, Outdated Libraries, Injection Risks, Logic Errors, and Performance Issues.')

doc.add_heading('2. Risk Level Calculation Rules', level=1)

doc.add_heading('2.1 Category Weights', level=2)
table1 = doc.add_table(rows=6, cols=2)
table1.style = 'Light Grid Accent 1'
header = table1.rows[0].cells
header[0].text = 'Category'
header[1].text = 'Weight Multiplier'
for cell in header:
    cell.paragraphs[0].runs[0].font.bold = True
data = [
    ('Security Flaws', '3.0x'),
    ('Injection Risks', '2.5x'),
    ('Outdated Libraries', '2.0x'),
    ('Logic Errors', '1.5x'),
    ('Performance Issues', '1.0x')
]
for i, (cat, weight) in enumerate(data, 1):
    row = table1.rows[i].cells
    row[0].text = cat
    row[1].text = weight
    row[0].paragraphs[0].runs[0].font.bold = True

doc.add_heading('2.2 Severity Scores', level=2)
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Light Grid Accent 1'
header2 = table2.rows[0].cells
header2[0].text = 'Severity'
header2[1].text = 'Points'
for cell in header2:
    cell.paragraphs[0].runs[0].font.bold = True
severities = [('CRITICAL', '10'), ('HIGH', '7'), ('MEDIUM', '4'), ('LOW', '1')]
for i, (sev, pts) in enumerate(severities, 1):
    row = table2.rows[i].cells
    row[0].text = sev
    row[1].text = pts
    row[0].paragraphs[0].runs[0].font.bold = True
    if sev == 'CRITICAL':
        row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)

doc.add_heading('2.3 Calculation Formula', level=2)
formula = doc.add_paragraph()
formula.add_run('For each finding: Contribution = Severity Points × Category Weight\n').bold = True
formula.add_run('Total Score = Sum of all contributions\n\n')
formula.add_run('Normalization (if >10 findings): Score × (10 ÷ Number of Findings)\n\n')
formula.add_run('Risk Level Thresholds:\n').bold = True
formula.add_run('• Score ≥ 60 → CRITICAL\n')
formula.add_run('• Score ≥ 40 → HIGH\n')
formula.add_run('• Score ≥ 20 → MEDIUM\n')
formula.add_run('• Score < 20 → LOW')

doc.add_heading('3. Score Conversion Rules', level=1)
table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Light Grid Accent 1'
header3 = table3.rows[0].cells
header3[0].text = 'Risk Level'
header3[1].text = 'Final Score (0-10)'
for cell in header3:
    cell.paragraphs[0].runs[0].font.bold = True
scores = [('CRITICAL', '1'), ('HIGH', '3'), ('MEDIUM', '5'), ('LOW', '8')]
for i, (risk, score) in enumerate(scores, 1):
    row = table3.rows[i].cells
    row[0].text = risk
    row[1].text = score
    row[0].paragraphs[0].runs[0].font.bold = True

doc.add_heading('4. Percentage Calculation Rules', level=1)
doc.add_paragraph('For each category, percentage = Sum of (Severity Weight × Count)')
table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Light Grid Accent 1'
header4 = table4.rows[0].cells
header4[0].text = 'Severity'
header4[1].text = 'Weight for %'
for cell in header4:
    cell.paragraphs[0].runs[0].font.bold = True
pct_weights = [('CRITICAL', '25'), ('HIGH', '15'), ('MEDIUM', '8'), ('LOW', '3')]
for i, (sev, weight) in enumerate(pct_weights, 1):
    row = table4.rows[i].cells
    row[0].text = sev
    row[1].text = weight
    row[0].paragraphs[0].runs[0].font.bold = True
doc.add_paragraph('Maximum: 100% (capped)').italic = True

doc.add_page_break()

# ============================================================================
# PART 2: COMPLETE TEST EXECUTION RESULTS
# ============================================================================
doc.add_heading('PART 2: COMPLETE TEST EXECUTION RESULTS', 0)

# Test 1
doc.add_heading('TEST 1/9: Hardcoded Credentials', level=1)
doc.add_paragraph('Description: Code with hardcoded password - should detect high/critical security flaw').italic = True
doc.add_paragraph('Expected Risk Level: HIGH | Actual: CRITICAL ✗').bold = True

doc.add_paragraph('Code to Analyze:').bold = True
code1 = """import os
import mysql.connector

# Database connection with hardcoded password
db_password = "mySecretPassword123"
connection = mysql.connector.connect(
    host="localhost",
    user="admin",
    password=db_password
)

def get_user_data(user_id):
    query = f"SELECT * FROM users WHERE id = {user_id}"
    cursor = connection.cursor()
    cursor.execute(query)
    return cursor.fetchall()"""
doc.add_paragraph(code1, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 1/10 | Risk Level: CRITICAL')
doc.add_paragraph('Total Findings: 4 (Security Flaws: 3, Injection Risks: 1)')

doc.add_paragraph('Security Flaws (3):').bold = True
findings1 = [
    ('hardcoded_password', 'HIGH', 'Line 6', 'Password is hardcoded in source code', 'db_password = \'mySecretPassword123\'', 'Move password to environment variable: password = os.getenv(\'DB_PASSWORD\')'),
    ('insecure_authentication', 'HIGH', 'Line 6', 'Insecure authentication due to hardcoded password', 'mysql.connector.connect(host=\'localhost\', user=\'admin\', password=db_password)', 'Implement a more secure authentication mechanism'),
    ('missing_input_validation', 'MEDIUM', 'Line 10', 'User input is not validated', 'def get_user_data(user_id):', 'Validate and sanitize user_id input')
]
for i, (typ, sev, line, desc, code, fix) in enumerate(findings1, 1):
    p = doc.add_paragraph()
    p.add_run(f'[{i}] Type: {typ}\n').bold = True
    p.add_run(f'Severity: {sev} | Line: {line}\n')
    p.add_run(f'Description: {desc}\n')
    p.add_run(f'Code: {code}\n')
    p.add_run(f'Fix: {fix}')

doc.add_paragraph('Injection Risks (1):').bold = True
p = doc.add_paragraph()
p.add_run('[1] Type: sql_injection\n').bold = True
p.add_run('Severity: HIGH | Line: 10\n')
p.add_run('Description: SQL injection risk due to unsanitized user input\n')
p.add_run('Code: query = f\'SELECT * FROM users WHERE id = {user_id}\'\n')
p.add_run('Fix: Use parameterized queries to avoid SQL injection')

doc.add_paragraph('Recommended Fixes:').bold = True
for i, fix in enumerate(['Move password to environment variable', 'Use parameterized queries', 'Implement secure authentication', 'Validate and sanitize user_id input'], 1):
    doc.add_paragraph(f'{i}. {fix}')

doc.add_paragraph()  # Spacing

# Test 2
doc.add_heading('TEST 2/9: SQL Injection Vulnerability', level=1)
doc.add_paragraph('Description: Code with SQL injection vulnerability').italic = True
doc.add_paragraph('Expected Risk Level: HIGH | Actual: MEDIUM ✗').bold = True

doc.add_paragraph('Code to Analyze:').bold = True
code2 = """def get_user(username):
    import sqlite3
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    query = "SELECT * FROM users WHERE username = '" + username + "'"
    cursor.execute(query)
    return cursor.fetchone()"""
doc.add_paragraph(code2, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 5/10 | Risk Level: MEDIUM')
doc.add_paragraph('Total Findings: 2 (Security Flaws: 1, Injection Risks: 1)')

doc.add_paragraph('Security Flaws (1):').bold = True
p = doc.add_paragraph()
p.add_run('[1] Type: missing_input_validation\n').bold = True
p.add_run('Severity: HIGH | Line: 6\n')
p.add_run('Description: User input (username) is not validated\n')
p.add_run('Fix: Validate and sanitize user inputs')

doc.add_paragraph('Injection Risks (1):').bold = True
p = doc.add_paragraph()
p.add_run('[1] Type: SQL_injection\n').bold = True
p.add_run('Severity: HIGH | Line: 6\n')
p.add_run('Description: SQL query vulnerable to injection\n')
p.add_run('Fix: Use parameterized queries: cursor.execute(\'SELECT * FROM users WHERE username = ?\', (username,))')

doc.add_paragraph()  # Spacing

# Test 3
doc.add_heading('TEST 3/9: Command Injection', level=1)
doc.add_paragraph('Description: Code with command injection risk').italic = True
doc.add_paragraph('Expected Risk Level: HIGH | Actual: MEDIUM ✗').bold = True

doc.add_paragraph('Code to Analyze:').bold = True
code3 = """import os
import subprocess

def ping_host(hostname):
    command = f"ping -c 4 {hostname}"
    result = subprocess.run(command, shell=True, capture_output=True)
    return result.stdout.decode()"""
doc.add_paragraph(code3, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 5/10 | Risk Level: MEDIUM')
doc.add_paragraph('Total Findings: 2 (Security Flaws: 1, Injection Risks: 1)')

doc.add_paragraph('Findings:').bold = True
doc.add_paragraph('1. Missing input validation (HIGH) - Validate hostname input')
doc.add_paragraph('2. Command injection (HIGH) - Use list and avoid shell=True')

doc.add_paragraph()  # Spacing

# Test 4
doc.add_heading('TEST 4/9: Outdated Library', level=1)
doc.add_paragraph('Description: Code using outdated library').italic = True
p = doc.add_paragraph()
p.add_run('Expected Risk Level: MEDIUM | Actual: MEDIUM ✓ PASSED').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph('Code to Analyze:').bold = True
code4 = """# Using requests 2.25.0 which has known vulnerabilities
import requests
# requests version: 2.25.0

response = requests.get("https://api.example.com/data")
data = response.json()
print(data)"""
doc.add_paragraph(code4, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 5/10 | Risk Level: MEDIUM')
doc.add_paragraph('Total Findings: 1 (Outdated Libraries: 1)')

doc.add_paragraph('Outdated Libraries (1):').bold = True
p = doc.add_paragraph()
p.add_run('[1] Type: vulnerable_package\n').bold = True
p.add_run('Severity: CRITICAL | Line: 1\n')
p.add_run('Description: Using requests==2.25.0 which has known CVE-2021-33503\n')
p.add_run('Fix: Update to requests>=2.28.0: pip install --upgrade requests')

doc.add_paragraph()  # Spacing

# Test 5
doc.add_heading('TEST 5/9: Logic Error - Off by One', level=1)
doc.add_paragraph('Description: Code with logic error').italic = True
p = doc.add_paragraph()
p.add_run('Expected Risk Level: LOW | Actual: LOW ✓ PASSED').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph('Code to Analyze:').bold = True
code5 = """def process_items(items):
    # Off-by-one error: should be len(items) instead of len(items) - 1
    for i in range(len(items) - 1):
        if items[i] > items[i + 1]:
            items[i], items[i + 1] = items[i + 1], items[i]
    return items"""
doc.add_paragraph(code5, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 8/10 | Risk Level: LOW')
doc.add_paragraph('Total Findings: 1 (Logic Errors: 1)')

doc.add_paragraph('Logic Errors (1):').bold = True
p = doc.add_paragraph()
p.add_run('[1] Type: off-by-one_error\n').bold = True
p.add_run('Severity: MEDIUM | Line: 3\n')
p.add_run('Description: Off-by-one error will cause last item to be excluded\n')
p.add_run('Fix: Change range to \'range(len(items))\' to include all items')

doc.add_paragraph()  # Spacing

# Test 6
doc.add_heading('TEST 6/9: Performance Issue - N+1 Query', level=1)
doc.add_paragraph('Description: Code with N+1 query problem').italic = True
doc.add_paragraph('Expected Risk Level: LOW | Actual: MEDIUM ✗').bold = True

doc.add_paragraph('Code to Analyze:').bold = True
code6 = """def get_user_posts(users):
    posts = []
    for user in users:
        # N+1 query problem
        user_posts = db.query(f"SELECT * FROM posts WHERE user_id = {user.id}")
        posts.extend(user_posts)
    return posts"""
doc.add_paragraph(code6, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 5/10 | Risk Level: MEDIUM')
doc.add_paragraph('Total Findings: 2 (Injection Risks: 1, Performance Issues: 1)')

doc.add_paragraph('Findings:').bold = True
doc.add_paragraph('1. SQL injection (HIGH) - Use parameterized queries')
doc.add_paragraph('2. N+1 query problem (MEDIUM) - Optimize with single query')

doc.add_paragraph()  # Spacing

# Test 7
doc.add_heading('TEST 7/9: Secure Code', level=1)
doc.add_paragraph('Description: Secure code example').italic = True
p = doc.add_paragraph()
p.add_run('Expected Risk Level: LOW | Actual: LOW ✓ PASSED').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph('Code to Analyze:').bold = True
code7 = """import os
from sqlalchemy import create_engine, text
import bcrypt

# Secure database connection using environment variables
db_password = os.getenv("DB_PASSWORD")
engine = create_engine(f"postgresql://user:{db_password}@localhost/db")

def get_user_data(user_id):
    # Parameterized query prevents SQL injection
    with engine.connect() as conn:
        result = conn.execute(text("SELECT * FROM users WHERE id = :id"), {"id": user_id})
        return result.fetchone()

def hash_password(password):
    # Secure password hashing
    salt = bcrypt.gensalt()
    return bcrypt.hashpw(password.encode(), salt)"""
doc.add_paragraph(code7, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
p = doc.add_paragraph()
p.add_run('Score: 8/10 | Risk Level: LOW\n').bold = True
p.add_run('Total Findings: 0\n\n')
p.add_run('✓ No vulnerabilities detected - code is secure!').bold = True
p.runs[2].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph()  # Spacing

# Test 8
doc.add_heading('TEST 8/9: Mixed Issues', level=1)
doc.add_paragraph('Description: Code with multiple vulnerability types').italic = True
doc.add_paragraph('Expected Risk Level: HIGH | Actual: CRITICAL ✗').bold = True

doc.add_paragraph('Code to Analyze:').bold = True
code8 = """import requests  # version 2.25.0 - outdated
import os

# Hardcoded API key
api_key = "sk-1234567890abcdef"
password = "admin123"

def fetch_data(user_input):
    # SQL injection risk
    query = f"SELECT * FROM data WHERE name = '{user_input}'"
    return query

# Command injection risk
def run_command(cmd):
    os.system(f"echo {cmd}")"""
doc.add_paragraph(code8, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 1/10 | Risk Level: CRITICAL')
doc.add_paragraph('Total Findings: 6 (Security Flaws: 3, Injection Risks: 2, Outdated Libraries: 1)')

doc.add_paragraph('Security Flaws (3):').bold = True
findings8 = [
    ('hardcoded_api_key', 'HIGH', 'Line 5', 'API key hardcoded', 'Move to environment variable'),
    ('hardcoded_password', 'HIGH', 'Line 6', 'Password hardcoded', 'Move to environment variable'),
    ('missing_input_validation', 'HIGH', 'Line 11', 'User input not validated', 'Validate using parameterized queries')
]
for i, (typ, sev, line, desc, fix) in enumerate(findings8, 1):
    p = doc.add_paragraph()
    p.add_run(f'[{i}] {typ} ({sev}) - {desc}\n').bold = True
    p.add_run(f'Fix: {fix}')

doc.add_paragraph('Injection Risks (2):').bold = True
doc.add_paragraph('[1] sql_injection (HIGH) - Use parameterized queries')
doc.add_paragraph('[2] command_injection (CRITICAL) - Use subprocess.run with list')

doc.add_paragraph('Outdated Libraries (1):').bold = True
doc.add_paragraph('[1] vulnerable_package (CRITICAL) - Update requests>=2.28.0')

doc.add_paragraph()  # Spacing

# Test 9
doc.add_heading('TEST 9/9: Eval/Exec Risk', level=1)
doc.add_paragraph('Description: Code using eval() - should detect code injection risk').italic = True
p = doc.add_paragraph()
p.add_run('Expected Risk Level: HIGH | Actual: HIGH ✓ PASSED').bold = True
p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph('Code to Analyze:').bold = True
code9 = """def calculate_expression(expression):
    # Dangerous use of eval
    result = eval(expression)
    return result

def execute_code(code_string):
    # Dangerous use of exec
    exec(code_string)"""
doc.add_paragraph(code9, style='Intense Quote')

doc.add_paragraph('Analysis Results:').bold = True
doc.add_paragraph('Score: 3/10 | Risk Level: HIGH')
doc.add_paragraph('Total Findings: 2 (Injection Risks: 2)')

doc.add_paragraph('Injection Risks (2):').bold = True
doc.add_paragraph('[1] code_injection (CRITICAL) - eval() allows arbitrary code execution')
doc.add_paragraph('[2] code_injection (CRITICAL) - exec() allows arbitrary code execution')
doc.add_paragraph('Fix: Implement safer expression parsing mechanism')

doc.add_page_break()

# Test Summary
doc.add_heading('TEST SUMMARY', level=1)
table_summary = doc.add_table(rows=10, cols=5)
table_summary.style = 'Light Grid Accent 1'
header_sum = table_summary.rows[0].cells
header_sum[0].text = 'Test'
header_sum[1].text = 'Status'
header_sum[2].text = 'Score'
header_sum[3].text = 'Risk Level'
header_sum[4].text = 'Findings'
for cell in header_sum:
    cell.paragraphs[0].runs[0].font.bold = True

summary_data = [
    ('Hardcoded Credentials', 'FAIL', '1/10', 'CRITICAL', '4'),
    ('SQL Injection', 'FAIL', '5/10', 'MEDIUM', '2'),
    ('Command Injection', 'FAIL', '5/10', 'MEDIUM', '2'),
    ('Outdated Library', 'PASS', '5/10', 'MEDIUM', '1'),
    ('Logic Error', 'PASS', '8/10', 'LOW', '1'),
    ('Performance Issue', 'FAIL', '5/10', 'MEDIUM', '2'),
    ('Secure Code', 'PASS', '8/10', 'LOW', '0'),
    ('Mixed Issues', 'FAIL', '1/10', 'CRITICAL', '6'),
    ('Eval/Exec Risk', 'PASS', '3/10', 'HIGH', '2'),
]

for i, (test, status, score, risk, findings) in enumerate(summary_data, 1):
    row = table_summary.rows[i].cells
    row[0].text = f'{i}. {test}'
    row[1].text = status
    row[2].text = score
    row[3].text = risk
    row[4].text = findings
    if status == 'PASS':
        row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
    else:
        row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 0, 0)
    row[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
summary_para = doc.add_paragraph()
summary_para.add_run('Overall Statistics:\n').bold = True
summary_para.add_run('• Total Tests: 9\n')
summary_para.add_run('• Passed: 4 (44.4%)\n')
summary_para.add_run('• Failed: 5 (55.6%)\n\n')
summary_para.add_run('Note: Failures are due to risk level mismatches (expected vs actual), but all vulnerabilities were correctly detected.').italic = True

# Save
output_path = project_root / "Code_Vulnerability_Complete_Report.docx"
try:
    doc.save(str(output_path))
    print(f"\n{'='*80}")
    print(f"   DOCX GENERATED SUCCESSFULLY")
    print(f"{'='*80}")
    print(f"   File: {output_path}")
    print(f"{'='*80}")
except PermissionError:
    import time
    timestamp = int(time.time())
    output_path = project_root / f"Code_Vulnerability_Complete_Report_{timestamp}.docx"
    doc.save(str(output_path))
    print(f"\n{'='*80}")
    print(f"   DOCX GENERATED (saved with timestamp)")
    print(f"{'='*80}")
    print(f"   File: {output_path}")
    print(f"{'='*80}")

